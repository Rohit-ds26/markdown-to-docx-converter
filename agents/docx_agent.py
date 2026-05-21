from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from markdown_it import MarkdownIt
from mdit_py_plugins.gfm import gfm_plugin

from agents.word_theme import (
    ACCENT,
    BODY,
    LINK,
    apply_premium_theme,
    ensure_reference_docx,
    polish_saved_docx,
    style_blockquote,
    style_code_block,
    style_heading_paragraph,
    style_horizontal_rule,
    style_table,
)


def _has_system_pandoc() -> bool:
    return shutil.which("pandoc") is not None


def _has_pypandoc() -> bool:
    try:
        import pypandoc  # noqa: F401

        return True
    except ImportError:
        return False


def _pandoc_extra_args() -> list[str]:
    ref = ensure_reference_docx()
    return [
        "--standalone",
        f"--reference-doc={ref}",
        "--syntax-highlighting=tango",
        "--toc",
        "--toc-depth=3",
        "--number-sections",
    ]


def _md_to_docx_pandoc_cli(*, markdown: str, output_path: str) -> None:
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as td:
        md_path = Path(td) / "input.md"
        md_path.write_text(markdown, encoding="utf-8")
        subprocess.run(
            [
                "pandoc",
                str(md_path),
                "-f",
                "gfm",
                "-t",
                "docx",
                "-o",
                str(out),
                *_pandoc_extra_args(),
            ],
            check=True,
            capture_output=True,
            text=True,
        )


def _md_to_docx_pypandoc(*, markdown: str, output_path: str) -> None:
    import pypandoc

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    pypandoc.convert_text(
        markdown,
        to="docx",
        format="gfm",
        outputfile=str(out),
        extra_args=_pandoc_extra_args(),
    )


def _render_inline(paragraph, inline_token) -> None:
    if not getattr(inline_token, "children", None):
        if inline_token.content:
            paragraph.add_run(inline_token.content)
        return

    bold = False
    italic = False

    for child in inline_token.children:
        t = child.type
        if t == "text":
            run = paragraph.add_run(child.content)
            run.bold = bold
            run.italic = italic
            run.font.color.rgb = BODY
        elif t == "strong_open":
            bold = True
        elif t == "strong_close":
            bold = False
        elif t == "em_open":
            italic = True
        elif t == "em_close":
            italic = False
        elif t == "code_inline":
            run = paragraph.add_run(child.content)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = ACCENT
        elif t in ("softbreak", "hardbreak"):
            paragraph.add_run("\n")
        elif t == "link_open":
            pass
        elif t == "link_close":
            pass
        elif t == "html_inline" and child.content:
            paragraph.add_run(child.content)
        elif t == "image":
            alt = child.attrs.get("alt", "") if child.attrs else ""
            if alt:
                run = paragraph.add_run(f"[{alt}]")
                run.italic = True
                run.font.color.rgb = LINK


def _add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    style_code_block(p)
    run = p.add_run(code.rstrip("\n"))
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)


def _add_blockquote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    style_blockquote(p)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = BODY


def _md_to_docx_builtin(*, markdown: str, output_path: str) -> None:
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    md = MarkdownIt("commonmark", {"html": False}).use(gfm_plugin)
    tokens = md.parse(markdown)
    doc = Document()
    apply_premium_theme(doc)

    list_stack: list[str] = []
    first_h1 = True
    i = 0
    while i < len(tokens):
        t = tokens[i]

        if t.type in ("fence", "code_block"):
            _add_code_block(doc, t.content)
            i += 1
            continue

        if t.type == "heading_open":
            level = min(int(t.tag[1]), 3)
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            h = doc.add_heading(level=level)
            if inline and inline.type == "inline":
                _render_inline(h, inline)
            is_first = level == 1 and first_h1
            if level == 1:
                first_h1 = False
            style_heading_paragraph(h, level, is_first_h1=is_first)
            i += 3
            continue

        if t.type == "bullet_list_open":
            list_stack.append("bullet")
            i += 1
            continue

        if t.type == "ordered_list_open":
            list_stack.append("ordered")
            i += 1
            continue

        if t.type in ("bullet_list_close", "ordered_list_close"):
            if list_stack:
                list_stack.pop()
            i += 1
            continue

        if t.type == "list_item_open":
            j = i + 1
            inline = None
            while j < len(tokens) and tokens[j].type != "list_item_close":
                if tokens[j].type == "inline":
                    inline = tokens[j]
                j += 1
            style = "List Bullet" if (list_stack and list_stack[-1] == "bullet") else "List Number"
            p = doc.add_paragraph(style=style)
            if inline:
                _render_inline(p, inline)
            i = j + 1
            continue

        if t.type == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            if inline and inline.type == "inline" and inline.content.strip():
                p = doc.add_paragraph()
                _render_inline(p, inline)
            i += 3
            continue

        if t.type == "blockquote_open":
            j = i + 1
            while j < len(tokens) and tokens[j].type != "blockquote_close":
                if tokens[j].type == "inline" and tokens[j].content.strip():
                    _add_blockquote(doc, tokens[j].content)
                j += 1
            i = j + 1
            continue

        if t.type == "hr":
            p = doc.add_paragraph()
            style_horizontal_rule(p)
            i += 1
            continue

        if t.type == "table_open":
            rows: list[list[str]] = []
            j = i + 1
            current_row: list[str] = []
            while j < len(tokens) and tokens[j].type != "table_close":
                if tokens[j].type == "tr_open":
                    current_row = []
                elif tokens[j].type == "inline":
                    current_row.append(tokens[j].content)
                elif tokens[j].type == "tr_close" and current_row:
                    rows.append(current_row)
                j += 1

            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                for r_idx, row in enumerate(rows):
                    for c_idx in range(cols):
                        table.rows[r_idx].cells[c_idx].text = row[c_idx] if c_idx < len(row) else ""
                style_table(table)
            i = j + 1
            continue

        i += 1

    doc.save(str(out))


def markdown_to_docx(*, markdown: str, output_path: str, engine: str = "auto") -> str:
    if engine not in {"auto", "pandoc", "builtin", "pypandoc"}:
        raise ValueError("engine must be one of: auto, pandoc, pypandoc, builtin")

    if engine in {"auto", "pypandoc"} and _has_pypandoc():
        try:
            _md_to_docx_pypandoc(markdown=markdown, output_path=output_path)
            return _finalize_docx(output_path, "pypandoc (premium template + TOC)")
        except Exception:
            if engine == "pypandoc":
                raise

    if engine in {"auto", "pandoc"} and _has_system_pandoc():
        _md_to_docx_pandoc_cli(markdown=markdown, output_path=output_path)
        return _finalize_docx(output_path, "pandoc (premium template + TOC)")

    if engine == "pandoc" and not _has_system_pandoc():
        raise RuntimeError(
            "Pandoc not found. Install from https://pandoc.org/ or: pip install pypandoc_binary"
        )

    _md_to_docx_builtin(markdown=markdown, output_path=output_path)
    return _finalize_docx(output_path, "builtin (premium styled)")


def _finalize_docx(output_path: str, engine_label: str) -> str:
    polish_saved_docx(output_path)
    return engine_label
