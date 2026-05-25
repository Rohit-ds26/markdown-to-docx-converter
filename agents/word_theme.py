"""
Premium Word styling — shared by Pandoc reference template and built-in converter.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, register_element_cls
from docx.oxml.ns import qn, nsmap
from docx.shared import Inches, Pt, RGBColor

# Register VML namespace for watermarks
try:
    from lxml import etree
    # VML namespace for watermark shapes
    nsmap['v'] = 'urn:schemas-microsoft-com:vml'
    nsmap['w10'] = 'urn:schemas-microsoft-com:office:word'
    nsmap['o'] = 'urn:schemas-microsoft-com:office:office'
except:
    pass

# Palette (modern report / SaaS document)
NAVY = RGBColor(0x0F, 0x17, 0x2A)
ACCENT = RGBColor(0x25, 0x63, 0xEB)
ACCENT_DARK = RGBColor(0x1E, 0x40, 0xAF)
MUTED = RGBColor(0x64, 0x74, 0x8B)
BODY = RGBColor(0x33, 0x41, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LINK = RGBColor(0x25, 0x63, 0xEB)

FILL_HEADER = "1E40AF"
FILL_ROW_ALT = "F8FAFC"
FILL_CODE = "F1F5F9"
FILL_QUOTE = "EFF6FF"
FILL_ACCENT_BAR = "2563EB"


def reference_docx_path() -> Path:
    return Path(__file__).resolve().parent.parent / "assets" / "reference.docx"


def trampolyne_logo_path() -> Path:
    return Path(__file__).resolve().parent.parent / "assets" / "trampolyne_logo.png"


def ensure_trampolyne_logo() -> Path:
    """
    Ensure the Trampolyne AI logo exists. If not, create a placeholder.
    This is automatically called when generating watermarks.
    """
    logo_path = trampolyne_logo_path()
    
    # If logo already exists, return it
    if logo_path.is_file() and logo_path.stat().st_size > 100:
        return logo_path
    
    # Create assets directory if it doesn't exist
    logo_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Try to create a nice logo with PIL
    try:
        from PIL import Image, ImageDraw, ImageFont
        
        # Create a new image with transparent background
        width, height = 800, 200
        img = Image.new('RGBA', (width, height), (0, 0, 0, 0))
        draw = ImageDraw.Draw(img)
        
        # Draw text "Trampolyne AI" in cyan/blue color
        text = "Trampolyne AI"
        
        # Try to use a nice font, fallback to default
        try:
            font = ImageFont.truetype("arial.ttf", 80)
        except:
            try:
                font = ImageFont.truetype("calibri.ttf", 80)
            except:
                font = ImageFont.load_default()
        
        # Calculate text position (centered)
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        x = (width - text_width) // 2
        y = (height - text_height) // 2
        
        # Draw text in cyan color (similar to Trampolyne AI branding)
        draw.text((x, y), text, fill=(37, 99, 235, 180), font=font)
        
        # Save the image
        img.save(str(logo_path))
        print("✓ Created placeholder Trampolyne AI logo")
        print("  Replace assets/trampolyne_logo.png with your actual logo for best results.")
        
    except ImportError:
        # PIL not available - create minimal PNG
        print("PIL/Pillow not installed. Creating minimal placeholder logo.")
        
        # Minimal PNG file (1x1 transparent pixel)
        png_data = (
            b'\x89PNG\r\n\x1a\n'  # PNG signature
            b'\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01'
            b'\x08\x06\x00\x00\x00\x1f\x15\xc4\x89'
            b'\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01'
            b'\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82'
        )
        
        with open(str(logo_path), 'wb') as f:
            f.write(png_data)
        
        print("✓ Created minimal placeholder. Please replace with actual Trampolyne AI logo.")
    
    return logo_path


def add_background_watermark(section, text: str = "Trampolyne AI", logo_path: Path | None = None) -> None:
    """
    Add a diagonal background watermark using Word's textpath shape.
    This creates a proper diagonal watermark that Word will render correctly.
    """
    # Skip watermark during initial creation - we'll add it in polish_saved_docx
    pass


def _set_paragraph_shading(paragraph, fill_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    p_pr.append(shd)


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_paragraph_bottom_border(paragraph, color_hex: str = "2563EB", size: int = 12) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _set_paragraph_left_border(paragraph, color_hex: str = "2563EB", size: int = 24) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def apply_premium_theme(doc: Document) -> None:
    """Page layout + style definitions for built-in conversion."""
    logo_path = ensure_trampolyne_logo()  # Automatically create logo if missing
    
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)
        
        # Add Trampolyne AI diagonal background watermark (applies to all pages)
        add_background_watermark(section, text="Trampolyne AI", logo_path=logo_path)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    pf = normal.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.22

    specs = (
        (1, 26, NAVY, Pt(18), Pt(10)),
        (2, 16, ACCENT_DARK, Pt(14), Pt(8)),
        (3, 13, MUTED, Pt(10), Pt(6)),
    )
    for level, size, color, before, after in specs:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri Light"
        h.font.bold = True
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.paragraph_format.space_before = before
        h.paragraph_format.space_after = after

    for style_name in ("List Bullet", "List Number"):
        if style_name in doc.styles:
            ls = doc.styles[style_name]
            ls.font.name = "Calibri"
            ls.font.size = Pt(11)
            ls.font.color.rgb = BODY
            ls.paragraph_format.space_after = Pt(4)


def style_heading_paragraph(paragraph, level: int, *, is_first_h1: bool = False) -> None:
    """Extra polish on heading paragraphs after content is added."""
    if level == 1:
        if is_first_h1:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_bottom_border(paragraph, color_hex="2563EB", size=16)
    elif level == 2:
        _set_paragraph_bottom_border(paragraph, color_hex="CBD5E1", size=6)


def style_code_block(paragraph) -> None:
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(10)
    _set_paragraph_shading(paragraph, FILL_CODE)
    _set_paragraph_left_border(paragraph, FILL_ACCENT_BAR, size=28)


def style_blockquote(paragraph) -> None:
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(8)
    _set_paragraph_shading(paragraph, FILL_QUOTE)
    _set_paragraph_left_border(paragraph, FILL_ACCENT_BAR, size=32)


def style_table(table) -> None:
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if r_idx == 0:
                _set_cell_shading(cell, FILL_HEADER)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = WHITE
                        run.font.name = "Calibri"
                        run.font.size = Pt(10.5)
            elif r_idx % 2 == 0:
                _set_cell_shading(cell, FILL_ROW_ALT)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)


def style_horizontal_rule(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(14)
    _set_paragraph_bottom_border(paragraph, color_hex="94A3B8", size=8)


def build_reference_docx(path: Path | None = None) -> Path:
    """
    Pandoc reference document with matching typography (Heading 1–3, Normal, etc.).
    """
    out = path or reference_docx_path()
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    apply_premium_theme(doc)

    # Pandoc maps markdown elements to these style names.
    if "Title" in doc.styles:
        t = doc.styles["Title"]
        t.font.name = "Calibri Light"
        t.font.size = Pt(28)
        t.font.bold = True
        t.font.color.rgb = NAVY

    if "First Paragraph" in doc.styles:
        fp = doc.styles["First Paragraph"]
        fp.font.name = "Calibri"
        fp.font.size = Pt(11)
        fp.font.color.rgb = BODY

    if "Block Text" in doc.styles:
        bt = doc.styles["Block Text"]
        bt.font.name = "Calibri"
        bt.font.italic = True
        bt.font.color.rgb = MUTED

    if "Source Code" in doc.styles:
        sc = doc.styles["Source Code"]
        sc.font.name = "Consolas"
        sc.font.size = Pt(9.5)

    # Seed sample content so Word retains style definitions reliably.
    h1 = doc.add_heading("Document Title", level=1)
    style_heading_paragraph(h1, 1, is_first_h1=True)
    doc.add_heading("Section", level=2)
    doc.add_paragraph("Body text sample.")
    doc.save(str(out))
    return out


def ensure_reference_docx() -> Path:
    path = reference_docx_path()
    if not path.is_file():
        build_reference_docx(path)
    return path


def polish_saved_docx(output_path: str) -> None:
    """
    Apply premium borders, table styling, margins, and DIAGONAL WATERMARK to an existing DOCX.
    """
    path = Path(output_path)
    if not path.is_file():
        return

    doc = Document(str(path))
    apply_premium_theme(doc)

    first_h1 = True
    for paragraph in doc.paragraphs:
        style_name = (paragraph.style.name if paragraph.style else "") or ""
        lower = style_name.lower()

        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", "").strip() or "1")
            except ValueError:
                level = 1
            is_first = level == 1 and first_h1
            if level == 1:
                first_h1 = False
            style_heading_paragraph(paragraph, level, is_first_h1=is_first)
        elif "quote" in lower or "block text" in lower:
            style_blockquote(paragraph)
        elif "source code" in lower or "verbatim" in lower or "code" in lower:
            style_code_block(paragraph)
            for run in paragraph.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)

    for table in doc.tables:
        style_table(table)

    # Add diagonal watermark to all sections
    logo_path = ensure_trampolyne_logo()  # Automatically create logo if missing
    for section in doc.sections:
        _add_diagonal_watermark_to_section(section, "Trampolyne AI", logo_path)

    doc.save(str(path))


def _add_diagonal_watermark_to_section(section, text: str, logo_path: Path | None = None) -> None:
    """
    Add diagonal watermark using direct XML manipulation with proper VML shapetype.
    """
    from lxml import etree
    
    header = section.header
    
    # Get first paragraph in header
    if not header.paragraphs:
        para = header.add_paragraph()
    else:
        para = header.paragraphs[0]
    
    # Create watermark using lxml with proper namespaces
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'v': 'urn:schemas-microsoft-com:vml',
        'o': 'urn:schemas-microsoft-com:office:office',
        'w10': 'urn:schemas-microsoft-com:office:word'
    }
    
    # Create run element
    r = etree.Element('{%s}r' % nsmap['w'])
    
    # Create pict element
    pict = etree.SubElement(r, '{%s}pict' % nsmap['w'])
    
    # First, add shapetype definition (required for textpath)
    shapetype = etree.SubElement(pict, '{%s}shapetype' % nsmap['v'])
    shapetype.set('id', '_x0000_t136')
    shapetype.set('coordsize', '21600,21600')
    shapetype.set('{%s}spt' % nsmap['o'], '136')
    shapetype.set('adj', '10800')
    shapetype.set('path', 'm@7,l@8,m@5,21600l@6,21600e')
    
    # Add formulas
    formulas = etree.SubElement(shapetype, '{%s}formulas' % nsmap['v'])
    for eqn in [
        'sum #0 0 10800', 'prod #0 2 1', 'sum 21600 0 @1', 'sum 0 0 @2',
        'sum 21600 0 @3', 'if @0 @3 0', 'if @0 21600 @1', 'if @0 0 @2',
        'if @0 @4 21600', 'mid @5 @6', 'mid @8 @5', 'mid @7 @8',
        'mid @6 @7', 'sum @6 0 @5'
    ]:
        f = etree.SubElement(formulas, '{%s}f' % nsmap['v'])
        f.set('eqn', eqn)
    
    # Add path
    path_elem = etree.SubElement(shapetype, '{%s}path' % nsmap['v'])
    path_elem.set('textpathok', 't')
    path_elem.set('{%s}connecttype' % nsmap['o'], 'custom')
    path_elem.set('{%s}connectlocs' % nsmap['o'], '@9,0;@10,10800;@11,21600;@12,10800')
    path_elem.set('{%s}connectangles' % nsmap['o'], '270,180,90,0')
    
    # Add textpath
    tp = etree.SubElement(shapetype, '{%s}textpath' % nsmap['v'])
    tp.set('on', 't')
    tp.set('fitshape', 't')
    
    # Add handles
    handles = etree.SubElement(shapetype, '{%s}handles' % nsmap['v'])
    h = etree.SubElement(handles, '{%s}h' % nsmap['v'])
    h.set('position', '#0,bottomRight')
    h.set('xrange', '6629,14971')
    
    # Add lock
    lock = etree.SubElement(shapetype, '{%s}lock' % nsmap['o'])
    lock.set('{%s}ext' % nsmap['v'], 'edit')
    lock.set('text', 't')
    lock.set('shapetype', 't')
    
    # Now create the actual shape instance
    shape = etree.SubElement(pict, '{%s}shape' % nsmap['v'])
    shape.set('id', 'PowerPlusWaterMarkObject')
    shape.set('type', '#_x0000_t136')
    shape.set('{%s}spid' % nsmap['o'], '_x0000_s2050')
    
    # CRITICAL: Use style attribute with rotation in fd units (-45 degrees = -2949120 fd)
    shape.set('style', 'position:absolute;margin-left:0;margin-top:0;width:527.85pt;height:131.95pt;rotation:-2949120fd;z-index:-251657216;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin')
    shape.set('{%s}allowincell' % nsmap['o'], 'f')
    shape.set('fillcolor', '#E0F0FF')  # Even lighter blue
    shape.set('stroked', 'f')
    
    # Add textpath element to shape
    textpath = etree.SubElement(shape, '{%s}textpath' % nsmap['v'])
    textpath.set('style', 'font-family:"Calibri";font-size:1pt;font-weight:bold')
    textpath.set('string', text)
    
    # Add wrap
    wrap = etree.SubElement(shape, '{%s}wrap' % nsmap['w10'])
    wrap.set('type', 'none')
    
    # Append to paragraph
    para._element.append(r)
    
    print(f"✓ Added diagonal watermark '{text}' with shapetype (rotation: -2949120fd)")
